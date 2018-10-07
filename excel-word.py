# -*- coding: utf-8 -*-
"""
Created on Sat Aug  4 18:52:35 2018

@author: lenovo
"""
import pandas as pd
import docx
from docx import Document
Data = pd.read_excel(r'bau.xlsx',\
                     header = 0, index_col = 0, sheet_name = [0])
Data = Data.T

list1 = []
list2 = []
list3 = []
list4 = []

for i in range(len(Data.columns)-1):
    list1.append(Data.iloc[0,i])
    list2.append(Data.iloc[1,i])
    list3.append(Data.iloc[2,i])
    list4.append(Data.iloc[3,i])
    
###更改不同获取模板文档路径
file=docx.Document('C:\\Users\\lenovo\\Desktop\\practice\\BAU简化信息模板.docx')
print("段落数:"+str(len(file.paragraphs))) #输出段落数
file_word = file
###更改不同保存路径
for i in range(len(list1)):
    path="C:\\Users\\lenovo\\Desktop\\practice\\search company product\\all company\\写入表格\\"+list1[i]+".docx"
    file_word.save(path)
    
    d = Document(path)
    t = d.tables[0]
    
    t.cell(1,1).text=str(list1[i])
    
    t.cell(2,3).text=str(list2[i])
    
    t.cell(1,3).text=str(list3[i])
    
    t.cell(2,1).text=str(list4[i])
    d.save(path)
